import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
import docx.oxml as oxml
import docx.oxml.ns as ns
import fg_parser
import word_renderer
import story_parser
import npc_parser
import item_parser
import vehicle_parser
import starship_parser
from ruleset_factory import RulesetFactory

def apply_safe_bookmarks(doc):
    print("Running document post-processor: Injecting live bookmark links...")
    return

def add_paragraph_bookmark(paragraph, bookmark_name, b_id):
    p_element = paragraph._p
    b_start = oxml.shared.OxmlElement('w:bookmarkStart')
    b_start.set(ns.qn('w:id'), str(b_id))
    b_start.set(ns.qn('w:name'), bookmark_name)
    b_end = oxml.shared.OxmlElement('w:bookmarkEnd')
    b_end.set(ns.qn('w:id'), str(b_id))
    p_element.insert(0, b_start)
    p_element.append(b_end)

def main(custom_mod_file=None, custom_output_file=None):
    try:
        mod_file = custom_mod_file or "FS_MtN.mod"
        output_file = custom_output_file or "Mission_to_Nagol_Campaign.docx"

        if not os.path.exists(mod_file):
            print(f"[!] Error: Could not find module package: '{mod_file}'")
            return f"[!] Error: Could not find module package: '{mod_file}'"

        print(f"--- Initiating Fantasy Grounds Module Extraction Pipeline ---")
        print(f"Target Archive: {mod_file}")

        mod_zip = zipfile.ZipFile(mod_file, 'r')
        
        display_name = ""
        try:
            with mod_zip.open('definition.xml') as def_file:
                def_tree = ET.parse(def_file)
                def_root = def_tree.getroot()
            name_node = def_root.find('name')
            if name_node is not None and name_node.text:
                display_name = name_node.text.strip()
            ruleset_node = def_root.find('ruleset')
            ruleset_name = ruleset_node.text.strip() if ruleset_node is not None else "Unknown"
        except KeyError:
            ruleset_name = "Unknown"

        print(f"[SYSTEM] Module metadata verified. Active Ruleset: '{ruleset_name}'")
        
        if custom_output_file == "__GET_INTERNAL_NAME__":
            mod_zip.close()
            return display_name if display_name else os.path.splitext(os.path.basename(mod_file))[0]

        rf = RulesetFactory(ruleset_name)

        try:
            xml_file = mod_zip.open('db.xml')
            db_root = ET.fromstring(xml_file.read())
            print("--- Successfully loaded the Fantasy Grounds XML data! ---")
        except KeyError:
            print("[!] Critical Error: db.xml was not found inside the module archive.")
            mod_zip.close()
            return

        story_template = rf.get_template_path("story")
        doc = Document(story_template)

        # Purge all template placeholder text cleanly
        for paragraph in list(doc.paragraphs):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Reference Tracking Registers for Post-Processing Hyperlinks
        processed_npcs = []
        processed_items = []
        processed_vehicles = []
        processed_starships = []
        processed_locations = []
        processed_tables = []
        processed_quests = []

        pipeline_order = [
            {"id": "story",      "is_appendix": False},
            {"id": "npc",        "is_appendix": True},
            {"id": "item",       "is_appendix": True},
            {"id": "vehicle",    "is_appendix": True},
            {"id": "starship",   "is_appendix": True},
            {"id": "location",   "is_appendix": True},
            {"id": "table",      "is_appendix": True},
            {"id": "quest",      "is_appendix": True}
        ]

        appendix_letters = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S", "T"]
        appendix_ptr = 0

        for step in pipeline_order:
            component = step["id"]
            current_appendix = f"Appendix {appendix_letters[appendix_ptr]}" if step["is_appendix"] else ""

            if component == "story":
                print(f"\nProcessing structured narrative hierarchy...")
                story_pages = story_parser.harvest_ordered_story_pages(db_root)
                for element in story_pages:
                    item_type = element[0]
                    if item_type == 'chapter':
                        doc.add_heading(element[1], level=1)
                    elif item_type == 'subchapter':
                        doc.add_heading(element[1], level=2)
                    elif item_type == 'page':
                        page_name, page_node = element[1], element[2]
                        doc.add_heading(page_name, level=3)
                        
                        blocks_node = page_node.find('blocks')
                        if blocks_node is not None:
                            ordered_blocks = []
                            for block in blocks_node:
                                order_val = block.find("order[@type='number']")
                                order_num = int(order_val.text) if order_val is not None and order_val.text else 999
                                ordered_blocks.append((order_num, block))
                            ordered_blocks.sort(key=lambda x: x[0])
                            
                            for order_num, block in ordered_blocks:
                                block_type = fg_parser.clean_xml_text(block.find("blocktype[@type='string']"))
                                if block_type == 'header':
                                    header_text = fg_parser.clean_xml_text(block.find("text[@type='string']"))
                                    if header_text:
                                        doc.add_paragraph(header_text, style=word_renderer.get_safe_style(doc, 'Heading 4'))
                                
                                # Process text blocks cleanly
                                text_node = block.find("text[@type='formattedtext']")
                                if text_node is not None:
                                    word_renderer.write_formatted_text(text_node, doc, block_type=block_type, xml_root=db_root, mod_zip=mod_zip)
                                    
                                # FIXED: Intercept image blocks directly at the block root level 
                                image_node = block.find(".//image") or block.find("image")
                                if image_node is not None:
                                    if block_type == "image" or block_type == "imageright" or block_type == "imageleft":
                                        # Pass the direct image wrapper to the formatter
                                        word_renderer.write_formatted_text(block, doc, block_type=block_type, xml_root=db_root, mod_zip=mod_zip)
                                    else:
                                        word_renderer.write_formatted_text(image_node, doc, block_type=block_type, xml_root=db_root, mod_zip=mod_zip)
            else:
                # DYNAMICAL ROUTING FOR ALL APPENDICES
                data_payload = rf.execute_parser(component, db_root)
                if data_payload:
                    renderer_module = rf.get_renderer_module(component)
                    if renderer_module:
                        
                        # FIXED FIX: Expanded type evaluation checks to catch items and flat dict lists safely
                        has_records = False
                        if isinstance(data_payload, list) and len(data_payload) > 0:
                            if isinstance(data_payload[0], dict):
                                if "npcs" in data_payload[0] or "starships" in data_payload[0]:
                                    inner_count = sum(len(group.get("npcs", []) or group.get("starships", [])) for group in data_payload)
                                    if inner_count > 0: has_records = True
                                else:
                                    has_records = True  # Flat lists of items/vehicles are valid records!
                            else:
                                has_records = True

                        if not has_records:
                            continue

                        if component == "npc":
                            for group in data_payload:
                                for n in group["npcs"]:
                                    processed_npcs.append((n["name"], f"REF_NPC_{n['raw_node'].tag}"))
                        elif component == "item":
                            for item in data_payload:
                                processed_items.append((item["name"], f"REF_ITEM_{item['raw_node'].tag}"))
                        elif component == "vehicle":
                            for veh in data_payload:
                                processed_vehicles.append((veh["name"], f"REF_VEHICLE_{veh['raw_node'].tag}"))
                        elif component == "starship":
                            for group in data_payload:
                                for ship in group["starships"]:
                                    processed_starships.append((ship["name"], f"REF_STARSHIP_{ship['raw_node'].tag}"))
                        elif component == "location":
                            for loc in data_payload:
                                processed_locations.append((loc["name"], f"REF_LOCATION_{loc['raw_node'].tag}"))
                        elif component == "table":
                            for t in data_payload:
                                processed_tables.append((t["name"], f"REF_TABLE_{t['raw_node'].tag}"))
                        elif component == "quest":
                            for q in data_payload:
                                processed_quests.append((q["name"], f"REF_QUEST_{q['raw_node'].tag}"))

                        template_file = rf.get_template_path(component)
                        render_func_name = f"render_{component}_appendix"
                        render_func = getattr(renderer_module, render_func_name, None)
                        
                        if render_func is None:
                            render_func_name = f"render_{component}s_appendix"
                            render_func = getattr(renderer_module, render_func_name, None)
                        
                        if render_func:
                            render_func(mod_zip, data_payload, doc, template_file, current_appendix)
                            if step["is_appendix"]: 
                                appendix_ptr += 1
                        else:
                            print(f"[!] Engine Error: Function '{render_func_name}' missed.")
                                
        print("\nResolving cross-reference anchors safely...")
        b_id = 2000
        all_targets = (processed_npcs + processed_items + processed_vehicles + 
                    processed_starships + processed_locations + processed_tables + processed_quests)
        
        for paragraph in doc.paragraphs:
            text_line = paragraph.text.strip()
            if not text_line: continue
            for name, bookmark_name in list(all_targets):
                if (text_line == name or text_line.startswith(name) or name in text_line or text_line.endswith("– " + name)):
                    add_paragraph_bookmark(paragraph, bookmark_name, b_id)
                    b_id += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text_line = paragraph.text.strip()
                        for name, bookmark_name in list(all_targets):
                            if text_line == name or name in text_line:
                                add_paragraph_bookmark(paragraph, bookmark_name, b_id)
                                b_id += 1

        mod_zip.close()
        doc.save(output_file)
        print("\n--- SUCCESS ---")
        print(f"Generated clean module workbook: {output_file}")
        return "SUCCESS"
        
    except Exception as e:
        import traceback
        print(f"\nProcessing Pipeline Failure: {e}")
        traceback.print_exc()
        return f"Failure: {str(e)}"

if __name__ == "__main__":
    main()