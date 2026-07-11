import os
import io
import docx.oxml as oxml
import docx.oxml.ns as ns
from docx.shared import Inches
from docxtpl import DocxTemplate, RichText
from PIL import Image as PILImage

def create_item_image_subdoc(tpl, mod_zip, asset_path, target_width=3.5):
    if not asset_path:
        return ""
    if "@" in asset_path:
        return f"[ASSET SOURCED EXTERNALLY: {asset_path}]"

    try:
        with mod_zip.open(asset_path) as img_file:
            img_data = img_file.read()
            
        with PILImage.open(io.BytesIO(img_data)) as img:
            width_px, height_px = img.size
            aspect_ratio = height_px / width_px
            target_height = target_width * aspect_ratio
            
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG")
            output_stream.seek(0)
            
        subdoc = tpl.new_subdoc()
        p = subdoc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(output_stream, width=Inches(target_width), height=Inches(target_height))
        return subdoc
    except KeyError:
        return f"[MISSING MODULE ASSET: {asset_path}]"
    except Exception as e:
        print(f"  [!] ITEM ASSET WARNING: Failed to process picture '{asset_path}'. Error: {e}")
        return f"[CORRUPT ASSET: {asset_path}]"

def render_item_appendix(mod_zip, items_data, master_doc, template_path, appendix_label=""):
    if not os.path.exists(template_path):
        print(f"[!] Error: Item renderer blueprint missing: '{template_path}'")
        return False

    print(f"Compiling {len(items_data)} items into text layout tables...")
    tpl = DocxTemplate(template_path)

    for item in items_data:
        raw_node = item["raw_node"]
        raw_id = raw_node.tag
        bookmark_name = f"REF_ITEM_{raw_id}"

        # --- INJECT INTERNAL LINK BOOKMARK ---
        bookmark_subdoc = tpl.new_subdoc()
        p_book = bookmark_subdoc.add_paragraph()
        
        b_start = oxml.shared.OxmlElement('w:bookmarkStart')
        b_start.set(ns.qn('w:id'), raw_id.replace("id-", ""))
        b_start.set(ns.qn('w:name'), bookmark_name)
        p_book._p.append(b_start)
        
        b_end = oxml.shared.OxmlElement('w:bookmarkEnd')
        b_end.set(ns.qn('w:id'), raw_id.replace("id-", ""))
        p_book._p.append(b_end)

        # Handle Image Asset Injection
        pic_path = raw_node.findtext("picture") or ""
        item["picture"] = create_item_image_subdoc(tpl, mod_zip, pic_path, target_width=3.5)

        # Handle RichText Descriptions
        desc_node = raw_node.find("item_description")
        if desc_node is not None and len(desc_node.findall('p')) > 0:
            rt = RichText()
            paragraphs = desc_node.findall('p')
            for idx, p_node in enumerate(paragraphs):
                text = p_node.text.strip() if p_node.text else ""
                if text:
                    rt.add(text)
                    if idx < len(paragraphs) - 1:
                        rt.add("\n\n") 
            item["item_description"] = rt
        else:
            item["item_description"] = "No description cataloged."

    context = {
        "items": items_data,
        "appendix_label": appendix_label
    }

    try:
        tpl.render(context)
        for element in tpl.element.body:
            master_doc.element.body.append(element)
        print("[SUCCESS] Equipment appendix compiled into chronicle workbook.")
        return True
    except Exception as render_err:
        print(f"[!] Item list single-pass layout parsing failed. Error Details: {render_err}")
        return False