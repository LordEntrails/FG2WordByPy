import os
import io
from docx.shared import Inches
from docxtpl import DocxTemplate
from PIL import Image as PILImage
from fg_parser import clean_xml_text
from word_renderer import write_formatted_text

def create_starship_image_subdoc(tpl, mod_zip, asset_path, target_width=3.5):
    if not asset_path or "@" in asset_path:
        return ""
    try:
        with mod_zip.open(asset_path) as img_file:
            img_data = img_file.read()
        with PILImage.open(io.BytesIO(img_data)) as img:
            width_px, height_px = img.size
            aspect_ratio = height_px / width_px
            target_height = target_width * aspect_ratio
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG")
            output_stream.seek(0)
            
        subdoc = tpl.new_subdoc()
        p = subdoc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(output_stream, width=Inches(target_width), height=Inches(target_height))
        return subdoc
    except Exception as e:
        print(f"  [TELEMETRY] Image subdoc build failed for {asset_path}: {e}")
        return ""

def render_starships_appendix(mod_zip, structured_categories, doc_base, blueprint_path, appendix_label=""):
    print("\n--- [TELEMETRY START: STARSHIPS RENDERER] ---")
    print(f"[DIAGNOSTIC] Blueprint Path: '{blueprint_path}'")
    # print(f"[DIAGNOSTIC] Appendix Label: '{appendix_label}'")
    
    if not os.path.exists(blueprint_path):
        print(f"[!] TELEMETRY ERROR: Starship template file does not exist at path!")
        return False

    if not structured_categories:
        print("[!] TELEMETRY WARNING: structured_categories array is completely EMPTY or None.")
        return False

    # print(f"[DIAGNOSTIC] Initializing DocxTemplate with blueprint.")
    tpl = DocxTemplate(blueprint_path)

    for c_idx, group in enumerate(structured_categories, 1):
        # print(f"  -> Category Folder [{c_idx}]: '{group.get('name')}' (Contains {len(group.get('starships', []))} starships)")
        
        for s_idx, ship in enumerate(group.get("starships", []), 1):
            raw_node = ship["raw_node"]
            # print(f"     [{s_idx}] Packaging Template Context variables for: '{ship['name']}'")
            
            # Blueprint Image Processing
            pic_path = raw_node.findtext("picture") or ""
            if pic_path:
                # print(f"       - Found picture node path: '{pic_path}'")
                ship["picture"] = create_starship_image_subdoc(tpl, mod_zip, pic_path, target_width=3.5)
            else:
                # print("       - No picture node path found.")
                ship["picture"] = ""
            
            # Description Processing
            desc_node = raw_node.find("description")
            if desc_node is not None:
                # print("       - Found description <formattedtext> block. Building subdoc...")
                subdoc_desc = tpl.new_subdoc()
                write_formatted_text(desc_node, subdoc_desc, xml_root=raw_node)
                ship["description"] = subdoc_desc
                # print(f"         Subdoc text generated paragraph count: {len(subdoc_desc.paragraphs)}")
            else:
                print("       - Description block missing.")
                ship["description"] = "No structural description logs archived."
                
            # Notes Processing
            notes_node = raw_node.find("notes")
            if notes_node is not None and (notes_node.text or list(notes_node)):
                print("       - Found notes <formattedtext> block. Building subdoc...")
                subdoc_notes = tpl.new_subdoc()
                write_formatted_text(notes_node, subdoc_notes, xml_root=raw_node)
                ship["ship_notes"] = subdoc_notes
            else:
                ship["ship_notes"] = None

    context = {
        "categories": structured_categories,
        "appendix_label": appendix_label
    }

    try:
        print("[DIAGNOSTIC] Executing tpl.render(context)...")
        tpl.render(context)
        
        # Investigate the template output body size
        tpl_elements = list(tpl.element.body)
        print(f"[DIAGNOSTIC] Render complete. Rendered Template Body contains {len(tpl_elements)} OXML elements.")
        
        if len(tpl_elements) <= 2:
            print("[!] CRITICAL ALERT: The template rendered almost nothing. Check your FS_Starships_Template.docx jinja tags!")
            print("    Ensure variables match: {% for cat in categories %} and {% for ship in cat.starships %}")

        print("[DIAGNOSTIC] Merging rendered elements into master document...")
        
        merge_count = 0
        # Filtered layout strategy to append elements and strip blank paragraph spaces safely using doc_base
        for element in tpl.element.body:
            if hasattr(element, 'tag'):
                if element.tag.endswith('sectPr'):
                    continue
                # Discard empty whitespace line breaks added by Word formatting blocks
                if element.tag.endswith('p') and not element.text and len(list(element)) == 0:
                    continue
            doc_base.element.body.append(element)
            merge_count += 1
            
        print(f"[SUCCESS] Merged {merge_count} elements into chronicle workbook tree.")
        print("--- [TELEMETRY END: STARSHIPS RENDERER] ---\n")
        return True
    except Exception as render_err:
        print(f"[!] CRITICAL FAILURE during template rendering execution: {render_err}")
        import traceback
        traceback.print_exc()
        return False