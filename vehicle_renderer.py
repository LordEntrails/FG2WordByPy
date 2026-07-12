import os
import io
from docx.shared import Inches
from docxtpl import DocxTemplate, RichText
from PIL import Image as PILImage

def create_vehicle_image_subdoc(tpl, mod_zip, asset_path, target_width):
    if not asset_path:
        return ""
    if "@" in asset_path:
        return f"[ASSET SOURCED EXTERNALLY: {asset_path}]"

    try:
        with mod_zip.open(asset_path) as img_file:
            img_data = img_file.read()
            
        with PILImage.open(io.BytesIO(img_data)) as img:
            width_px, height_px = img.size
            aspect_ratio = height_px / width_px
            target_height = target_width * aspect_ratio
            
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG")
            output_stream.seek(0)
            
        subdoc = tpl.new_subdoc()
        p = subdoc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(output_stream, width=Inches(target_width), height=Inches(target_height))
        return subdoc
    except KeyError:
        return f"[MISSING MODULE ASSET: {asset_path}]"
    except Exception as e:
        print(f"  [!] VEHICLE ASSET WARNING: Failed to process picture '{asset_path}'. Error: {e}")
        return f"[CORRUPT ASSET: {asset_path}]"

def render_vehicle_appendix(mod_zip, vehicles_data, master_doc, template_path, appendix_label=""):
    if not os.path.exists(template_path):
        print(f"[!] Error: Vehicle renderer blueprint missing: '{template_path}'")
        return False

    print(f"Compiling {len(vehicles_data)} vehicles into narrative spec sheets...")
    tpl = DocxTemplate(template_path)

    for vehicle in vehicles_data:
        raw = vehicle["raw_node"]
        
        # Image extraction passes
        pic_path = raw.findtext("picture") or ""
        token_flat = raw.findtext("token") or ""
        token_cam = raw.findtext("token3Dflat") or ""
        
        vehicle["picture"] = create_vehicle_image_subdoc(tpl, mod_zip, pic_path, target_width=3.5)
        vehicle["token_flat"] = create_vehicle_image_subdoc(tpl, mod_zip, token_flat, target_width=1.7)
        vehicle["token_camera"] = create_vehicle_image_subdoc(tpl, mod_zip, token_cam, target_width=1.7)

        # Extract description text (Standardized to .description context key)
        notes_node = raw.find("notes")
        if notes_node is not None and len(notes_node.findall('p')) > 0:
            rt = RichText()
            paragraphs = notes_node.findall('p')
            for idx, p_node in enumerate(paragraphs):
                text = p_node.text.strip() if p_node.text else ""
                if text:
                    rt.add(text)
                    if idx < len(paragraphs) - 1:
                        rt.add("\n\n")
            vehicle["description"] = rt
        else:
            vehicle["description"] = "No structural description logs archived."

    context = {
        "vehicles": vehicles_data,
        "appendix_label": appendix_label
    }

    try:
        tpl.render(context)
        
        # Filtered layout strategy to append elements and strip blank paragraph spaces
        for element in tpl.element.body:
            if hasattr(element, 'tag'):
                if element.tag.endswith('sectPr'):
                    continue
                # Discard empty whitespace line breaks added by Word formatting blocks
                if element.tag.endswith('p') and not element.text and len(list(element)) == 0:
                    continue
            master_doc.element.body.append(element)
        print("[SUCCESS] Vehicle appendix compiled into chronicle workbook.")
        return True
    except Exception as render_err:
        print(f"[!] Vehicle list layout parsing failed. Error Details: {render_err}")
        return False