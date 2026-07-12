import os
import io
from docx.shared import Inches
from docxtpl import DocxTemplate
from PIL import Image as PILImage
from fg_parser import clean_xml_text
from word_renderer import write_formatted_text, get_safe_style

def create_location_map_subdoc(tpl, mod_zip, xml_root):
    link_node = xml_root.find(".//link[@class='imagewindow']")
    if link_node is None:
        return ""
    record_target = link_node.get("recordname") or link_node.text or ""
    if not record_target:
        return ""
    image_id = record_target.split('.')[1] if '.' in record_target else record_target
    img_entry = xml_root.find(f"../../image/{image_id}/image/layers/layer/bitmap")
    if img_entry is None:
        img_entry = xml_root.find(f".//image/{image_id}/image/layers/layer/bitmap")
    if img_entry is None or not img_entry.text:
        return ""

    image_path = img_entry.text.strip()
    try:
        with mod_zip.open(image_path) as img_file:
            img_data = img_file.read()
        with PILImage.open(io.BytesIO(img_data)) as img:
            width_px, height_px = img.size
            aspect_ratio = height_px / width_px
            target_width = 3.5 
            target_height = target_width * aspect_ratio
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG")
            output_stream.seek(0)
            
        subdoc = tpl.new_subdoc()
        p = subdoc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(output_stream, width=Inches(target_width), height=Inches(target_height))
        return subdoc
    except Exception:
        return ""

def render_location_appendix(mod_zip, locations_data, master_doc, template_path, appendix_label=""):
    if not os.path.exists(template_path):
        print(f"[!] Error: Location template path missing: '{template_path}'")
        return False
        
    tpl = DocxTemplate(template_path)
    for loc in locations_data:
        raw = loc["raw_node"]
        loc["picture"] = create_location_map_subdoc(tpl, mod_zip, raw)
        
        subloc_node = raw.find("sublocations")
        if subloc_node is not None and list(subloc_node):
            sub_doc_sl = tpl.new_subdoc()
            for child_sl in subloc_node.findall('*'):
                p = sub_doc_sl.add_paragraph(style='List Paragraph')
                p.add_run("▪\t").bold = True
                p.add_run(child_sl.tag + " - Reference Tracking Vector Entry")
            loc["sublocation"] = sub_doc_sl
        else:
            loc["sublocation"] = "No nested operational sub-sectors cataloged."

        text_node = raw.find("text")
        if text_node is not None:
            subdoc_text = tpl.new_subdoc()
            write_formatted_text(text_node, subdoc_text, xml_root=raw)
            loc["location_text"] = subdoc_text
        else:
            loc["location_text"] = "No geographic survey logs archived."

        gm_node = raw.find("gmnotes")
        if gm_node is not None and (gm_node.text or list(gm_node)):
            subdoc_gm = tpl.new_subdoc()
            write_formatted_text(gm_node, subdoc_gm, xml_root=raw)
            loc["gm_notes"] = subdoc_gm
        else:
            loc["gm_notes"] = "No secondary master tracking data vectors registered."

    context = {"locations": locations_data, "appendix_label": appendix_label}
    tpl.render(context)
    master_doc.add_page_break()
    for element in tpl.element.body:
        master_doc.element.body.append(element)
    return True