import os
from docx.shared import Inches
from docxtpl import DocxTemplate
from fg_parser import clean_xml_text
from word_renderer import get_safe_style, insert_internal_hyperlink

def build_rollable_matrix_subdoc(tpl, doc_base, table_node):
    """
    Generates a sub-document grid containing formatted rollable table entries.
    Locks widths precisely cell-by-cell to fit a 3.5" total column margin.
    """
    subdoc = tpl.new_subdoc()
    try:
        col_count = int(clean_xml_text(table_node.find("resultscols")) or 1)
    except ValueError:
        col_count = 1
        
    total_cols = col_count + 1
    word_table = subdoc.add_table(rows=1, cols=total_cols)
    word_table.style = 'Grid Table 4 Accent 1'
    word_table.header_row = True
    word_table.row_banding = True
    
    # Grid column geometries factoring a 3.5" maximum width ceiling
    RANGE_WIDTH = Inches(0.65)
    DATA_WIDTH = Inches(2.85 / col_count)
    
    hdr_cells = word_table.rows[0].cells
    hdr_cells[0].text = "Range"
    hdr_cells[0].width = RANGE_WIDTH
    
    for c in range(1, total_cols):
        lbl_node = table_node.find(f"labelcol{c}")
        hdr_cells[c].text = clean_xml_text(lbl_node) if lbl_node is not None else f"Result {c}"
        hdr_cells[c].width = DATA_WIDTH

    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.style = get_safe_style(doc_base, 'Normal')
            for run in p.runs:
                run.bold = True

    rows_node = table_node.find("tablerows")
    if rows_node is not None:
        for row_node in rows_node.findall('*'):
            from_r = clean_xml_text(row_node.find("fromrange")) or ""
            to_r = clean_xml_text(row_node.find("torange")) or ""
            range_display = f"{from_r}-{to_r}" if from_r != to_r else from_r
            
            row_cells = word_table.add_row().cells
            row_cells[0].text = range_display
            row_cells[0].width = RANGE_WIDTH
            
            results_node = row_node.find("results")
            if results_node is not None:
                for r_idx, res_node in enumerate(results_node.findall('*'), 1):
                    if r_idx >= total_cols:
                        break
                    res_text = clean_xml_text(res_node.find("result")) or ""
                    cell_obj = row_cells[r_idx]
                    cell_obj.width = DATA_WIDTH
                    
                    p_res = cell_obj.paragraphs[0]
                    p_res.style = get_safe_style(doc_base, 'Normal')
                    
                    link_node = res_node.find("resultlink")
                    bookmark_anchor = ""
                    if link_node is not None:
                        c_type = link_node.findtext("class") or ""
                        r_ref = link_node.findtext("recordname") or ""
                        if c_type == "item" and "item.id-" in r_ref:
                            bookmark_anchor = f"REF_ITEM_{r_ref.split('.')[-1]}"
                            
                    if bookmark_anchor:
                        insert_internal_hyperlink(p_res, res_text, bookmark_anchor)
                    else:
                        p_res.add_run(res_text)

    p_space = subdoc.add_paragraph()
    p_space.paragraph_format.space_after = Inches(0.1)
    return subdoc

def render_tables_appendix(mod_zip, tables_data, master_doc, template_path, appendix_label=""):
    if not os.path.exists(template_path):
        print(f"[!] Error: Table template path missing: '{template_path}'")
        return False
        
    tpl = DocxTemplate(template_path)
    for item in tables_data:
        item["render_table_block"] = build_rollable_matrix_subdoc(tpl, master_doc, item["raw_node"])
        
    context = {"tables": tables_data, "appendix_label": appendix_label}
    tpl.render(context)
    master_doc.add_page_break()
    for element in tpl.element.body:
        master_doc.element.body.append(element)
    return True