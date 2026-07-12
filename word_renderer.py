import io
import docx.oxml as oxml
import docx.oxml.ns as ns
from docx import Document
from docx.shared import Inches
from PIL import Image
from fg_parser import clean_xml_text
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor

def get_safe_style(doc, style_name, default_fallback='Normal'):
    if style_name in doc.styles:
        return style_name
    return default_fallback

def apply_inline_styles(node, paragraph_obj, bold=False, italic=False, underline=False):
    """
    Recursively steps through text nodes, passing down formatting states
    so nested inline tags (e.g., bold + underline) inherit styles correctly.
    """
    if node is None:
        return

    # 1. Inherit historical parent flags or trigger new ones based on the current tag
    current_bold = bold or (node.tag == 'b')
    current_italic = italic or (node.tag == 'i')
    current_underline = underline or (node.tag == 'u')

    # 2. Apply accumulated formatting parameters to the text segment
    if node.text:
        cleaned_text = node.text.lstrip('\n\t').rstrip('\n\t')
        if cleaned_text or node.text == " ":
            run = paragraph_obj.add_run(cleaned_text)
            if current_bold: 
                run.bold = True
            if current_italic: 
                run.italic = True
            if current_underline: 
                run.underline = True

    # 3. Recurse into child nodes carrying forward the active style environment
    for child in node:
        apply_inline_styles(child, paragraph_obj, current_bold, current_italic, current_underline)
        
        # 4. Process tail text using the original parent context scope
        if child.tail:
            cleaned_tail = child.tail.lstrip('\n\t').rstrip('\n\t')
            if cleaned_tail or child.tail == " ":
                run = paragraph_obj.add_run(cleaned_tail)
                if bold: 
                    run.bold = True
                if italic: 
                    run.italic = True
                if underline: 
                    run.underline = True

def extract_and_insert_image(mod_zip, image_path, doc, target_width=3.5):
    try:
        with mod_zip.open(image_path) as img_file:
            image_stream = io.BytesIO(img_file.read())
            
        with Image.open(image_stream) as img:
            width_px, height_px = img.size
            aspect_ratio = height_px / width_px
            target_height = target_width * aspect_ratio
            
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG")
            output_stream.seek(0)
            
            p_img = doc.add_paragraph()
            p_img.alignment = 1  # Center aligned
            p_img.paragraph_format.space_before = Inches(0.1)
            p_img.paragraph_format.space_after = Inches(0.1)
            
            p_img.add_run().add_picture(output_stream, width=Inches(target_width), height=Inches(target_height))
            return True
            
    except KeyError:
        p = doc.add_paragraph(style=get_safe_style(doc, 'Normal'))
        p.add_run(f"[MISSING IMAGE ASSET: {image_path}]").font.color.rgb = RGBColor(200, 0, 0)
        return False
    except Exception as e:
        print(f"  [!] STORY IMAGE FAILURE: Failed to parse '{image_path}'. Error: {e}")
        return False

def insert_internal_hyperlink(paragraph, text, bookmark_name):
    """
    Inserts a native Microsoft Word cross-reference hyperlink into a paragraph's
    run collection without corrupting underlying table cell element structures.
    """
    run = paragraph.add_run()
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('w:anchor'), bookmark_name)
    
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')
    
    color = oxml.shared.OxmlElement('w:color')
    color.set(ns.qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = oxml.shared.OxmlElement('w:u')
    u.set(ns.qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    text_node = oxml.shared.OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    
    run._r.getparent().insert(run._r.getparent().index(run._r), hyperlink)
    return hyperlink

def embed_encounter_table(doc, record_name, xml_root):
    if xml_root is None or not record_name:
        return False
        
    node_id = record_name.split('.')[-1] if '.' in record_name else record_name
    battle_node = xml_root.find(f"./battle/{node_id}")
    
    if battle_node is None:
        return False

    battle_name = clean_xml_text(battle_node.find("name")) or "Unnamed Skirmish"
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Inches(0.15)
    p_title.paragraph_format.space_after = Inches(0.05)
    p_title.add_run(f"Battle Deployment: {battle_name}").bold = True
    
    npclist = battle_node.find("npclist")
    if npclist is None or not list(npclist):
        p = doc.add_paragraph()
        p.add_run("   (No combatants cataloged in this encounter structure)").italic = True
        return True

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Grid Table 4 Accent 1'
    table.header_row = True
    table.row_banding = True
    table.first_column = False
    table.last_column = False
    table.column_banding = False
    
    QTY_WIDTH = Inches(0.65)
    DESC_WIDTH = Inches(2.85)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Combatant Designation"
    hdr_cells[0].width = QTY_WIDTH
    hdr_cells[1].width = DESC_WIDTH
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.style = get_safe_style(doc, 'Normal')
            for run in p.runs:
                run.bold = True

    # Tracking variable to ensure unique local target IDs inside table frames
    local_b_id = 5000 

    for child in npclist.findall('*'):
        count = clean_xml_text(child.find("count")) or "1"
        npc_name = clean_xml_text(child.find("name")) or "Unknown Combatant"
        
        npc_link_node = child.find("link")
        bookmark_anchor = ""
        
        if npc_link_node is not None:
            rec_ref = npc_link_node.get("recordname") or npc_link_node.text or ""
            if "npc.id-" in rec_ref and "@" not in rec_ref:
                target_id = rec_ref.split('.')[-1]
                bookmark_anchor = f"REF_NPC_{target_id}"

        row_cells = table.add_row().cells
        row_cells[0].width = QTY_WIDTH
        row_cells[1].width = DESC_WIDTH
        row_cells[0].text = f"{count}x"
        
        p_name = row_cells[1].paragraphs[0]
        p_name.style = get_safe_style(doc, 'Normal')
        
        if bookmark_anchor:
            # 1. Create the live hyperlinked text run inside the cell
            insert_internal_hyperlink(p_name, npc_name, bookmark_anchor)
            
            # 2. FIXED ROOT CAUSE: Force-inject the OXML paragraph bookmark target definitions
            # directly onto the cell context so Word registers them immediately!
            import docx.oxml as oxml
            import docx.oxml.ns as ns
            p_element = p_name._p
            b_start = oxml.shared.OxmlElement('w:bookmarkStart')
            b_start.set(ns.qn('w:id'), str(local_b_id))
            b_start.set(ns.qn('w:name'), bookmark_anchor)
            b_end = oxml.shared.OxmlElement('w:bookmarkEnd')
            b_end.set(ns.qn('w:id'), str(local_b_id))
            
            p_element.insert(0, b_start)
            p_element.append(b_end)
            local_b_id += 1
        else:
            p_name.add_run(npc_name)
            
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Inches(0.1)
    return True

def embed_parcel_table(doc, record_name, xml_root):
    if xml_root is None or not record_name:
        return False
        
    node_id = record_name.split('.')[-1] if '.' in record_name else record_name
    parcel_node = xml_root.find(f"./treasureparcels/{node_id}")
    
    if parcel_node is None:
        return False

    parcel_name = clean_xml_text(parcel_node.find("name")) or "Unnamed Inventory Cache"
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Inches(0.15)
    p_title.paragraph_format.space_after = Inches(0.05)
    p_title.add_run(f"📦 Inventory Cache: {parcel_name}").bold = True
    
    coinlist = parcel_node.find("coinlist")
    itemlist = parcel_node.find("itemlist")
    
    has_coins = coinlist is not None and list(coinlist)
    has_items = itemlist is not None and list(itemlist)
    
    if not has_coins and not has_items:
        p = doc.add_paragraph()
        p.add_run("   (This cache contains no items or currency entries)").italic = True
        return True

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Grid Table 4 Accent 1'
    table.header_row = True
    table.row_banding = True
    table.first_column = False
    table.last_column = False
    table.column_banding = False
    
    QTY_WIDTH = Inches(0.65)
    DESC_WIDTH = Inches(2.85)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Item Description / Currency Type"
    hdr_cells[0].width = QTY_WIDTH
    hdr_cells[1].width = DESC_WIDTH
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.style = get_safe_style(doc, 'Normal')
            for run in p.runs:
                run.bold = True

    if has_coins:
        for coin in coinlist.findall('*'):
            amount = clean_xml_text(coin.find("amount")) or "0"
            denom = clean_xml_text(coin.find("description")) or "Credits"
            
            row_cells = table.add_row().cells
            row_cells[0].width = QTY_WIDTH
            row_cells[1].width = DESC_WIDTH
            row_cells[0].text = amount
            
            p_desc = row_cells[1].paragraphs[0]
            p_desc.style = get_safe_style(doc, 'Normal')
            p_desc.add_run(denom).italic = True

    if has_items:
        for item in itemlist.findall('*'):
            count = clean_xml_text(item.find("count")) or "1"
            item_name = clean_xml_text(item.find("name")) or "Unknown Item"
            
            bookmark_anchor = ""
            item_link_node = item.find("item_link") or item.find("link")
            if item_link_node is not None:
                rec_ref = item_link_node.get("recordname") or item_link_node.text or ""
                if "item.id-" in rec_ref:
                    target_id = rec_ref.split('.')[-1]
                    bookmark_anchor = f"REF_ITEM_{target_id}"

            row_cells = table.add_row().cells
            row_cells[0].width = QTY_WIDTH
            row_cells[1].width = DESC_WIDTH
            row_cells[0].text = f"{count}x"
            
            p_desc = row_cells[1].paragraphs[0]
            p_desc.style = get_safe_style(doc, 'Normal')
            
            if bookmark_anchor:
                insert_internal_hyperlink(p_desc, item_name, bookmark_anchor)
            else:
                p_desc.add_run(item_name)
                
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Inches(0.1)
    return True

def write_formatted_text(xml_node, doc, block_type=None, xml_root=None, mod_zip=None):
    """
    Interprets and renders formatted text, headers, links, and inline images.
    """
    if xml_node is None:
        return
    for child in xml_node:
        if child.tag == 'p':
            if block_type == 'frame' or xml_node.tag == 'frame':
                style = get_safe_style(doc, 'Chat Paragraph', 'Normal')
                p = doc.add_paragraph(style=style)
                apply_inline_styles(child, p)
            else:
                p = doc.add_paragraph(style=get_safe_style(doc, 'Normal'))
                apply_inline_styles(child, p)
                
        elif child.tag == 'h':
            p = doc.add_paragraph(style=get_safe_style(doc, 'Heading 5'))
            p.add_run(clean_xml_text(child))
            
        elif child.tag == 'frame':
            style = get_safe_style(doc, 'Chat Paragraph', 'Normal')
            p = doc.add_paragraph(style=style)
            apply_inline_styles(child, p)
            
        elif child.tag == 'li':
            p = doc.add_paragraph(style=get_safe_style(doc, 'List Paragraph'))
            p.add_run("•\t")
            apply_inline_styles(child, p)

        elif child.tag == 'image':
            image_path = child.findtext("image") or ""
            if image_path and mod_zip is not None:
                extract_and_insert_image(mod_zip, image_path, doc)
                
                caption_text = clean_xml_text(child.find("caption"))
                if caption_text:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = 1 
                    p_cap.paragraph_format.space_before = Inches(0.02)
                    p_cap.paragraph_format.space_after = Inches(0.1)
                    run_cap = p_cap.add_run(caption_text)
                    run_cap.italic = True

        elif child.tag in ['linklist', 'link']:
            links = child.findall('link') if child.tag == 'linklist' else [child]
            for link in links:
                link_class = link.get("class") or ""
                record_name = link.get("recordname") or link.text or ""
                link_text = clean_xml_text(link)
                
                if link_class == "battle":
                    success = embed_encounter_table(doc, record_name, xml_root)
                    if success:
                        continue
                        
                if link_class == "treasureparcel":
                    success = embed_parcel_table(doc, record_name, xml_root)
                    if success:
                        continue
                
                style = get_safe_style(doc, 'Internal Link', 'Normal')
                p = doc.add_paragraph(style=style)
                
                bookmark_name = ""
                # Parse internal anchor bookmarks only if it's a native campaign asset (no '@' symbol)
                if "@" not in record_name:
                    if "npc.id-" in record_name:
                        bookmark_name = f"REF_NPC_{record_name.split('.')[-1]}"
                    elif "item.id-" in record_name:
                        bookmark_name = f"REF_ITEM_{record_name.split('.')[-1]}"
                    elif "vehicle.id-" in record_name:
                        bookmark_name = f"REF_VEHICLE_{record_name.split('.')[-1]}"
                    elif "location.id-" in record_name:
                        bookmark_name = f"REF_LOCATION_{record_name.split('.')[-1]}"
                    elif "starships.id-" in record_name or "starship.id-" in record_name:
                        bookmark_name = f"REF_STARSHIP_{record_name.split('.')[-1]}"
                
                if bookmark_name:
                    insert_internal_hyperlink(p, link_text, bookmark_name)
                else:
                    # Write the baseline unlinked reference text run
                    p.add_run(link_text)
                    
                    # SYSTEM CHRONICLE EXTENSION: Capture external library module tracking nodes
                    if "@" in record_name:
                        external_source = record_name.split("@")[-1].replace("_", " ").strip()
                        if external_source:
                            # Add spacing and append the source text run styled in clean italics
                            run_source = p.add_run(f" (See: {external_source})")
                            run_source.italic = True

        elif child.tag == 'table':
            xml_rows = child.findall('tr')
            if xml_rows:
                max_cols = max((len(r.findall('td')) + len(r.findall('th'))) for r in xml_rows)
                word_table = doc.add_table(rows=len(xml_rows), cols=max_cols)
                
                word_table.style = 'Grid Table 4 Accent 1'
                word_table.header_row = True
                word_table.row_banding = True
                word_table.first_column = False
                word_table.last_column = False
                word_table.column_banding = False
                
                for r_idx, xml_row in enumerate(xml_rows):
                    cells = xml_row.findall('td') + xml_row.findall('th')
                    for c_idx, cell in enumerate(cells):
                        if c_idx >= max_cols:
                            continue
                            
                        word_cell = word_table.cell(r_idx, c_idx)
                        word_cell.text = clean_xml_text(cell)
                        
                        for paragraph in word_cell.paragraphs:
                            paragraph.style = get_safe_style(doc, 'Normal')
                            if r_idx == 0:
                                for run in paragraph.runs:
                                    run.bold = True